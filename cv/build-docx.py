"""Generate ATS-safe .docx versions of the CVs from the .html sources.

The HTML files are the single source of truth. This script parses them and
emits a matching Word document with the same single-column, no-table,
no-image structure that ATS parsers handle reliably.

Usage:
    python build-docx.py
"""

from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

HERE = Path(__file__).parent
FONT = "Calibri"

# Block kinds we care about, keyed by the tag/class they come from.
NAME, HEADLINE, CONTACT, SECTION, ENTRY, META, BODY, BULLET = (
    "name", "headline", "contact", "section", "entry", "meta", "body", "bullet",
)


class CVParser(HTMLParser):
    """Turn the CV HTML into a flat list of (kind, runs) blocks.

    A run is (text, bold, italic). A run of ("\n", ...) means a line break.
    """

    BLOCK_TAGS = {"h1", "h2", "h3", "p", "li"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks = []
        self.kind = None
        self.runs = []
        self.bold = 0
        self.italic = 0
        self.in_style = False
        self.in_skills = 0

    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)
        classes = attrs.get("class", "").split()

        if tag == "style":
            self.in_style = True
            return
        if tag == "div" and "skills" in classes:
            self.in_skills += 1
            return
        if tag == "br":
            if self.kind:
                self.runs.append(("\n", False, False))
            return
        if tag in ("strong", "b"):
            self.bold += 1
            return
        if tag in ("em", "i"):
            self.italic += 1
            return

        if tag in self.BLOCK_TAGS:
            self._flush()
            if tag == "h1":
                self.kind = NAME
            elif tag == "h2":
                self.kind = SECTION
            elif tag == "h3":
                self.kind = ENTRY
            elif tag == "li":
                self.kind = BULLET
            elif "headline" in classes:
                self.kind = HEADLINE
            elif "contact" in classes:
                self.kind = CONTACT
            elif "meta" in classes:
                self.kind = META
            else:
                self.kind = BODY

    def handle_endtag(self, tag):
        if tag == "style":
            self.in_style = False
        elif tag == "div" and self.in_skills:
            self.in_skills -= 1
        elif tag in ("strong", "b"):
            self.bold = max(0, self.bold - 1)
        elif tag in ("em", "i"):
            self.italic = max(0, self.italic - 1)
        elif tag in self.BLOCK_TAGS:
            self._flush()

    def handle_data(self, data):
        if self.in_style or not self.kind:
            return
        # Collapse the source's pretty-printing whitespace into single spaces.
        text = " ".join(data.split())
        if not text:
            return
        if self.runs and not self.runs[-1][0].endswith(("\n", " ")):
            text = " " + text
        self.runs.append((text, bool(self.bold), bool(self.italic)))

    def _flush(self):
        if self.kind and self.runs:
            self.blocks.append((self.kind, self.runs))
        self.kind, self.runs = None, []

    def close(self):
        super().close()
        self._flush()


def add_bottom_border(paragraph):
    """Section headings get a rule under them, as in the HTML."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    p_pr.append(borders)


# kind -> (size_pt, bold, italic, space_before_pt, space_after_pt)
# Sizes track the .html <style> block so the Word output paginates like the PDF.
STYLES = {
    NAME:     (18, True,  False, 0,  2),
    HEADLINE: (11, True,  False, 0,  2),
    CONTACT:  (9, False, False, 0,  10),
    SECTION:  (10.5, True, False, 11, 4),
    ENTRY:    (10, True,  False, 6,  0),
    META:     (9, False, True,  0,  2),
    BODY:     (10, False, False, 0, 4),
    BULLET:   (10, False, False, 0, 2),
}


def build(html_path: Path, docx_path: Path):
    parser = CVParser()
    parser.feed(html_path.read_text(encoding="utf-8"))
    parser.close()

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(43)

    for kind, runs in parser.blocks:
        size, bold, italic, before, after = STYLES[kind]

        if kind == BULLET:
            para = doc.add_paragraph(style="List Bullet")
        else:
            para = doc.add_paragraph()

        fmt = para.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.keep_with_next = kind in (SECTION, ENTRY)

        for text, run_bold, run_italic in runs:
            if text == "\n":
                para.add_run().add_break()
                continue
            run = para.add_run(text.upper() if kind == SECTION else text)
            run.font.name = FONT
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = bold or run_bold
            run.italic = italic or run_italic

        if kind == SECTION:
            add_bottom_border(para)

    doc.save(docx_path)
    print(f"  {docx_path.name}  ({len(parser.blocks)} blocks)")


def main():
    sources = sorted(HERE.glob("felipe-roque-*.html"))
    if not sources:
        raise SystemExit("No CV .html files found next to this script.")
    print("Building DOCX:")
    for html_path in sources:
        build(html_path, html_path.with_suffix(".docx"))


if __name__ == "__main__":
    main()
